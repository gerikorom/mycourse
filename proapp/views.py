from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse, JsonResponse, Http404, HttpResponseRedirect
from .models import *
from django.contrib.auth.decorators import login_required
from django.urls import reverse
from django.views.generic import CreateView, DeleteView, DetailView, UpdateView, ListView
from django.contrib.auth.mixins import LoginRequiredMixin
import pandas as pd
from .forms import *
import xlsxwriter as xlsx
from docx import Document
#pip install python-docx
from easy_pdf.rendering import render_to_pdf
#pip install django-easy-pdf
from django.core.mail import send_mail, EmailMultiAlternatives


from django.contrib.auth import login, authenticate

# Create your views here.

def kezdolap(request):
    SQL = Konyv.objects.raw('SELECT * FROM proapp_Konyv')
    konyvek = Konyv.objects.all().order_by('-id')[:4]
    elso = Konyv.objects.first()
    idalapjan = Konyv.objects.get(iro='Jókai Mór')
    gardonyi = Konyv.objects.filter(iro='Gárdonyi Géza')

    abcsorrend = Konyv.objects.all().order_by('iro')

    if request.method == 'POST' and request.is_ajax:
        form = ResztvetelForm(request.POST)
        if form.is_valid():
            if not Resztvetel.objects.filter(olvasoest=form.cleaned_data.get('olvasoest'), felhasznalo=request.user).exists():
                kivalasztottOlvasoEst = form.cleaned_data.get('olvasoest')
                if kivalasztottOlvasoEst.helyekszama != 0:
                    mentes = form.save(commit=False)
                    mentes.felhasznalo = request.user
                    mentes.save()
                    kivalasztottOlvasoEst.helyekszama = kivalasztottOlvasoEst.helyekszama - 1
                    kivalasztottOlvasoEst.save()

                    return JsonResponse({"uzenet":'Sikeres foglalás!'}, status=200)
                else:
                    return JsonResponse({"uzenet": 'Sajnos már nincs hely!'}, status=200)
            else:
                return JsonResponse({"uzenet": 'Már feliratkozott!'}, status=200)
        else:
            return JsonResponse({"error": form.errors.as_json()}, status=400)

    c = {
        'konyvek':konyvek,
        'form': ResztvetelForm
    }

    response = render(request, 'kezdolap.html', c)

    latogatasSzama = int(request.COOKIES.get('latogatas', '0'))
    response.set_cookie('latogatas', latogatasSzama+1)
    response.set_cookie('gyumolcs', 'alma')

    request.session['zoldseg'] = 'répa'
    request.session['kosar'] = ['Zöld poló', 'Piros zokni']

    return response

    #return HttpResponse(abcsorrend)
    #return render(request, 'kezdolap.html', c)

def regisztracio(request):
    if request.method == 'POST':
        form = SignUpForm(request.POST)
        if form.is_valid():
            form.save()
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password1')
            user = authenticate(username=username, password=password, request=request)
            login(request, user)
            return redirect('proapp:kezdolap')
    else:
        form = SignUpForm()

    c = {
        'form': form
    }
    return render(request, 'regisztracio.html', c)

def zoldseg(request):
    return HttpResponse(request.session.get('kosar', 'ures'))

@login_required
def konyvek(request):

    konyvek = Konyv.objects.all()

    utolsoId = request.session.get('konyvid', 0)

    c = {
        'object_list': konyvek,
        'utolsoId': utolsoId
    }

    return render(request, 'konyvek.html', c)

@login_required
def konyvfelvetel(request):

    form = KonyvForm()
    if request.method == 'POST':
        form = KonyvForm(request.POST)
        if form.is_valid():
            form.save()
            form = KonyvForm()
    '''
    form = FelvetelForm()
    if request.method == 'POST':
        form = FelvetelForm(request.POST)
        if form.is_valid():
            print(form.cleaned_data)
            Konyv.objects.create(**form.cleaned_data)
            form = FelvetelForm()
        else:
            print(form.errors)
    '''

    c = {
        'form':form
    }

    return render(request, 'konyvfelvetel.html', c)

@login_required
def kolcsonzes(request):

    form = KolcsonzesForm()
    if request.method == 'POST' and request.is_ajax:
        form = KolcsonzesForm(request.POST)
        if form.is_valid():
            mentes = form.save(commit=False)
            mentes.felhasznalo = request.user
            mentes.save()

            return JsonResponse({"adat": mentes.konyv.iro + ' ' + mentes.konyv.cim, 'ar': mentes.konyv.ar}, status=200)
        else:
            print(form.errors.items())
            for m, i in form.errors.items():
                print(m)
                for j in i:
                    print(j)

            return JsonResponse({"error": form.errors.as_json()}, status=400)

    c = {
        'form':form
    }

    return render(request, 'kolcsonzes.html', c)

@login_required
def estek(request):
    estek = OlvasoEst.objects.all()

    if request.method == 'POST' and request.is_ajax:
        form = ResztvetelForm(request.POST)
        if form.is_valid():
            if not Resztvetel.objects.filter(olvasoest=form.cleaned_data.get('olvasoest'), felhasznalo=request.user).exists():
                kivalasztottOlvasoEst = form.cleaned_data.get('olvasoest')
                if kivalasztottOlvasoEst.helyekszama != 0:
                    mentes = form.save(commit=False)
                    mentes.felhasznalo = request.user
                    mentes.save()
                    kivalasztottOlvasoEst.helyekszama = kivalasztottOlvasoEst.helyekszama - 1
                    kivalasztottOlvasoEst.save()

                    return JsonResponse({"uzenet":'Sikeres foglalás!'}, status=200)
                else:
                    return JsonResponse({"uzenet": 'Sajnos már nincs hely!'}, status=200)
            else:
                return JsonResponse({"uzenet": 'Már feliratkozott!'}, status=200)
        else:
           return JsonResponse({"error": form.errors.as_json()}, status=400)
    c = {
        'estek':estek,
        'form': ResztvetelForm
    }

    return render(request, 'estek.html', c)

@login_required
def konyv(request, id):
    #elem = Konyv.objects.get(id=id)

    '''
    elem = get_object_or_404(Konyv, id=id)
    '''
    try:
        elem = Konyv.objects.get(id=id)
        request.session['konyvid'] = id


        elem.delete()


    except Konyv.DoesNotExist:
        raise Http404


    c = {
        'object':elem
    }

    return render(request, 'konyv.html', c)

@login_required
def konyvar(request, ar):

    konyvek = Konyv.objects.filter(ar=ar)
    if not konyvek:
        return HttpResponse('Nincs ilyen könyv!')
    else:
        return HttpResponse(konyvek)


@login_required
def book_upload_view(request):
    form = ExcelFeltoltesForm()
    if request.method == "POST":
        form = ExcelFeltoltesForm(request.POST, request.FILES)
        if form.is_valid():
            excelFile = request.FILES['file']
            df = pd.read_excel(excelFile, engine='openpyxl')
            #pandas, openpyxl, xlrd?
            df.fillna('', inplace=True)
            print(df)
            '''
            for i in df.itertuples():
                Konyv.objects.create(cim=i[1],
                                     iro=i[2],
                                     leiras=i[3],
                                     ar=i[4],
                                     fajlnev=i[5])
            
            lista = [Konyv(cim=i[1],
                                    iro=i[2],
                                    leiras=i[3],
                                    ar=i[4],
                                    fajlnev=i[5]) for i in df.itertuples()]
            '''
            lista = []
            for i in df.itertuples():
                if not Konyv.objects.filter(cim=i[1], iro=i[2]).exists():
                    lista.append(Konyv(cim=i[1],
                                    iro=i[2],
                                    leiras=i[3],
                                    ar=i[4],
                                    fajlnev=i[5]))

            Konyv.objects.bulk_create(lista)


    c = {
        'form': form
    }

    return render(request, 'konyvfeltoltes.html', c)

@login_required
def excel_lekeres(request):
    if request.method == 'POST':
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename=konyvek.xlsx'
        workbook = xlsx.Workbook(response, {'in_memory': True})
        worksheet = workbook.add_worksheet('Könyv lista')

        datum = workbook.add_format({'num_format': 'yyyy.mm.dd', 'align': 'center'})
        kozep = workbook.add_format({'align': 'center', 'font_color':'FF0000', 'bold': True, 'italic':True})
        szam = workbook.add_format({'num_format': '#,##0.00'})

        worksheet.write(0, 0, 'Író', kozep)
        worksheet.write(0, 1, 'Cím', kozep)
        worksheet.write(0, 2, 'Ár', kozep)
        worksheet.write(0, 3, 'Leírás', kozep)

        lista = Konyv.objects.all()
        for sor, i in enumerate(lista, 1):
            worksheet.write(sor, 0, i.iro, kozep)
            worksheet.write(sor, 1, i.cim, kozep)
            worksheet.write(sor, 2, i.ar, szam)
            worksheet.write(sor, 3, i.leiras)


        workbook.close()
        return response

    return redirect('proapp:kezdolap')

@login_required
def word_lekeres(request):
    if request.method == "POST":
        document = Document()
        document.add_heading('Konyvek', 0)

        p = document.add_paragraph('Sima paragrafus ')
        p.add_run('félkövér').bold = True
        p.add_run(' és ')
        p.add_run('dőlt betűkkel.').italic = True

        document.add_paragraph('Idézet', style='Intense Quote')

        document.add_paragraph(
            'Rendezetlen lista', style='List Bullet'
        )
        document.add_paragraph(
            'Rendezett lista', style='List Number'
        )

        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Író'
        hdr_cells[1].text = 'Cím'
        hdr_cells[2].text = 'Ár'
        hdr_cells[3].text = 'Leírás'
        lista = Konyv.objects.all()
        for sor, i in enumerate(lista, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = i.iro
            row_cells[1].text = i.cim
            row_cells[2].text = str(i.ar)
            row_cells[3].text = i.leiras

        document.add_page_break()


        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=konyvek.docx'
        document.save(response)

        return response

    return redirect('proapp:kezdolap')

@login_required
def pdf_lekeres(request):
    if request.method == "POST":
        template = 'pdflekeres.html'
        lista = Konyv.objects.all()

        c = {
            'lista': lista
            }
        tt = render_to_pdf(template, c, content_type='application/pdf', response_class=HttpResponse)
        response = HttpResponse(tt, content_type='application/pdf')
        response['Content-Disposition'] = 'inline; filename=konyvek.pdf'

        send_mail(
            'Tárgy',
            'Üzenet',
            'proapp.proba@gmail.com',
            ['proapp.proba@gmail.com'],
            fail_silently=False,
        )


        subject, from_email, to = 'Tárgy', 'proapp.proba@gmail.com', 'proapp.proba@gmail.com'
        text_content = 'Ez egy fontos üzenet'
        html_content = '<p>Ez egy <strong>fontos</strong> üzenet.</p>'
        msg = EmailMultiAlternatives(subject, text_content, from_email, [to])
        msg.attach_alternative(html_content, "text/html")
        msg.attach('konyvek.pdf', tt, 'application/pdf')
        msg.send()

        return response

    return redirect('proapp:kezdolap')

class KonyvListView(LoginRequiredMixin, ListView):
    template_name = 'konyvek.html'
    queryset = Konyv.objects.all()

class KonyvDetailView(LoginRequiredMixin, DetailView):
    template_name = 'konyv.html'
    queryset = Konyv.objects.all()

class KonyvCreateView(LoginRequiredMixin, CreateView):
    template_name = 'konyvhozzaadas.html'
    form_class = KonyvForm
    queryset = Konyv.objects.all()



